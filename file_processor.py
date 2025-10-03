"""
File processing utilities for various document formats
Handles text extraction from PDF, DOCX, Excel, and other formats
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional, Union
import pandas as pd
from docx import Document
import PyPDF2
from config import Config

# Set up logging
logger = logging.getLogger(__name__)


class FileProcessor:
    """Handles processing of various file formats for PII detection"""
    
    def __init__(self):
        self.supported_extensions = Config.ALLOWED_EXTENSIONS
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file extension is supported"""
        if not filename:
            return False
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.supported_extensions
    
    def get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        if not self.is_supported_file(filename):
            return 'unsupported'
        return filename.rsplit('.', 1)[1].lower()
    
    def process_file(self, file_path: str, filename: str = None) -> Dict[str, Any]:
        """
        Process file and extract text content
        
        Args:
            file_path: Path to the file
            filename: Original filename (optional)
            
        Returns:
            Dictionary with extracted content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        filename = filename or os.path.basename(file_path)
        file_type = self.get_file_type(filename)
        
        if not self.is_supported_file(filename):
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            # Route to appropriate processor
            if file_type == 'txt':
                return self._process_text_file(file_path, filename)
            elif file_type == 'pdf':
                return self._process_pdf_file(file_path, filename)
            elif file_type == 'docx':
                return self._process_docx_file(file_path, filename)
            elif file_type in ['xlsx', 'csv']:
                return self._process_excel_file(file_path, filename)
            elif file_type == 'json':
                return self._process_json_file(file_path, filename)
            else:
                raise ValueError(f"Processor not implemented for: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
    
    def _process_text_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return {
                'filename': filename,
                'file_type': 'text',
                'content': content,
                'metadata': {
                    'size': len(content),
                    'lines': len(content.split('\n')),
                    'encoding': 'utf-8'
                }
            }
        except Exception as e:
            logger.error(f"Error reading text file {filename}: {str(e)}")
            raise
    
    def _process_pdf_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            content = ""
            metadata = {'pages': 0, 'extracted_pages': 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text
                            metadata['extracted_pages'] += 1
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {str(e)}")
                        continue
            
            return {
                'filename': filename,
                'file_type': 'pdf',
                'content': content.strip(),
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Error reading PDF file {filename}: {str(e)}")
            raise
    
    def _process_docx_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            doc = Document(file_path)
            content = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                content += f"\n--- Table {table_count + 1} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
                table_count += 1
            
            return {
                'filename': filename,
                'file_type': 'docx',
                'content': content.strip(),
                'metadata': {
                    'paragraphs': paragraph_count,
                    'tables': table_count,
                    'size': len(content)
                }
            }
        except Exception as e:
            logger.error(f"Error reading DOCX file {filename}: {str(e)}")
            raise
    
    def _process_excel_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process Excel/CSV file"""
        try:
            file_type = self.get_file_type(filename)
            
            if file_type == 'csv':
                df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
                sheets = {'Sheet1': df}
            else:
                # Excel file
                sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            content = ""
            metadata = {'sheets': len(sheets), 'total_rows': 0, 'total_columns': 0}
            
            for sheet_name, df in sheets.items():
                if len(sheets) > 1:
                    content += f"\n--- Sheet: {sheet_name} ---\n"
                
                # Add column headers
                content += " | ".join(df.columns.astype(str)) + "\n"
                content += "-" * (len(" | ".join(df.columns.astype(str)))) + "\n"
                
                # Add data rows
                for _, row in df.iterrows():
                    row_text = []
                    for value in row.values:
                        if pd.isna(value):
                            row_text.append("")
                        else:
                            row_text.append(str(value))
                    content += " | ".join(row_text) + "\n"
                
                metadata['total_rows'] += len(df)
                metadata['total_columns'] = max(metadata['total_columns'], len(df.columns))
            
            return {
                'filename': filename,
                'file_type': file_type,
                'content': content.strip(),
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Error reading Excel/CSV file {filename}: {str(e)}")
            raise
    
    def _process_json_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            content = self._json_to_text(data)
            
            return {
                'filename': filename,
                'file_type': 'json',
                'content': content,
                'metadata': {
                    'size': len(content),
                    'structure': self._analyze_json_structure(data)
                }
            }
        except Exception as e:
            logger.error(f"Error reading JSON file {filename}: {str(e)}")
            raise
    
    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """Convert JSON data to readable text format"""
        text = ""
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text += f"{prefix}{key}:\n"
                    text += self._json_to_text(value, prefix + "  ")
                else:
                    text += f"{prefix}{key}: {value}\n"
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text += f"{prefix}[{i}]:\n"
                text += self._json_to_text(item, prefix + "  ")
        else:
            text += f"{prefix}{data}\n"
        
        return text
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure for metadata"""
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': len(data),
                'nested_objects': sum(1 for v in data.values() if isinstance(v, dict)),
                'arrays': sum(1 for v in data.values() if isinstance(v, list))
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data))
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100] + ('...' if len(str(data)) > 100 else '')
            }
    
    def cleanup_file(self, file_path: str) -> bool:
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up file: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {str(e)}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            stat = os.stat(file_path)
            return {
                'size': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'modified': stat.st_mtime,
                'is_readable': os.access(file_path, os.R_OK)
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {}


# Global processor instance
_processor_instance = None

def get_file_processor() -> FileProcessor:
    """Get or create global file processor instance"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = FileProcessor()
    return _processor_instance